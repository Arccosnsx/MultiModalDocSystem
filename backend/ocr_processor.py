# ocr_processor.py
import os
from typing import Tuple
import docx
from PyPDF2 import PdfReader


class OcrProcessor:
    """OCR处理器，处理文件解析和文本提取"""
    TEXT_FILE_EXTENSIONS = {
        '.txt', '.md', '.text', '.log', '.csv', '.json',
        '.xml', '.yaml', '.yml', '.ini', '.cfg', '.py',
        '.js', '.html', '.css', '.java', '.cpp', '.c', '.h'
    }

    @staticmethod
    def extract_text_from_file(file_path: str) -> Tuple[str, bool]:
        """
        从文件中提取文本内容

        Args:
            file_path: 文件路径

        Returns:
            Tuple[str, bool]: (提取的文本内容, 是否需要OCR)
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension in OcrProcessor.TEXT_FILE_EXTENSIONS:
            return OcrProcessor._extract_text_file(file_path), False

        elif file_extension == '.docx':
            return OcrProcessor._extract_docx_file(file_path), False

        elif file_extension == '.pdf':
            return OcrProcessor._extract_pdf_file(file_path)
        else:
            return "", True

    @staticmethod
    def _extract_text_file(file_path: str) -> str:
        """提取文本文件内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            encodings = ['gbk', 'gb2312', 'latin-1', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return ""

    @staticmethod
    def _extract_docx_file(file_path: str) -> str:
        """提取Word文档内容"""
        try:
            doc = docx.Document(file_path)
            paragraphs = [
                para.text for para in doc.paragraphs if para.text.strip()]

            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))

            all_text = paragraphs + tables_text
            return '\n'.join(all_text)
        except Exception:
            return ""

    @staticmethod
    def _extract_pdf_file(file_path: str) -> Tuple[str, bool]:
        """提取PDF文件内容，返回是否需要OCR"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            extracted_text = text.strip()

            if extracted_text and len(extracted_text) > 50:
                return extracted_text, False
            else:
                return "", True
        except Exception:
            return "", True
